"""Build Ayan Anees's one-page ATS resume from the canonical public profile."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PROFILE_PATH = ROOT / "src" / "content" / "public-profile.json"
OUTPUT_PATH = ROOT / "docs" / "Ayan-Anees-Resume.draft.docx"

INK = RGBColor(20, 18, 15)
MUTED = RGBColor(77, 73, 68)
ACCENT = RGBColor(31, 85, 96)
RULE = "B8B1A7"
FONT = "Calibri"


def set_font(run, size: float, *, bold: bool = False, color=INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_spacing(paragraph, *, before=0, after=0, line=1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "271")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)
    set_spacing(paragraph, after=2, line=1.0)
    set_font(paragraph.add_run(text), 9.25)


def add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    set_spacing(paragraph, before=5, after=4, line=1.0)
    add_bottom_border(paragraph)
    set_font(paragraph.add_run(title.upper()), 10.5, bold=True, color=ACCENT)


def add_role_line(document: Document, left: str, right: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_TAB_ALIGNMENT.RIGHT)
    set_spacing(paragraph, after=1, line=1.0)
    set_font(paragraph.add_run(left), 10.25, bold=True)
    set_font(paragraph.add_run(f"\t{right}"), 9.4, color=MUTED)


def build() -> Path:
    profile = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Named ATS one-page override to compact_reference_guide.
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    document.core_properties.title = f"{profile['name']} - Resume"
    document.core_properties.subject = "ATS resume"
    document.core_properties.creator = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = "compact_reference_guide with ATS one-page and memo_masthead contact overrides"

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(name, after=1)
    set_font(name.add_run(profile["name"]), 22, bold=True)

    role = document.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(role, after=3)
    set_font(role.add_run(profile["role"]), 11.25, bold=True, color=ACCENT)

    contacts = profile["contacts"]
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=2)
    contact_text = f"{profile['location']} | {contacts['phone']} | {contacts['email']}"
    set_font(contact.add_run(contact_text), 9.25, color=MUTED)

    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(links, after=5)
    set_font(links.add_run(f"{profile['canonicalUrl']} | {contacts['github']}"), 9.0, color=MUTED)
    add_bottom_border(links)

    summary = (
        "Full-stack and AI developer building dependable web interfaces, document-processing workflows, "
        "and human-supervised AI systems. Experience spans React and TypeScript frontends, Python/FastAPI "
        "services, Azure document extraction, Databricks-backed assistance, and deterministic analytics."
    )
    add_section(document, "Profile")
    paragraph = document.add_paragraph()
    set_spacing(paragraph, after=2, line=1.02)
    set_font(paragraph.add_run(summary), 9.5)

    experience = profile["experience"][0]
    add_section(document, "Experience")
    add_role_line(document, f"{experience['organization']} | {experience['role']}", "June 2026 - Present | Lahore, Pakistan")
    bullet_id = create_bullet_numbering(document)
    bullets = [
        "Develop an internal AI operations and document-processing platform using Python, FastAPI, React, and SQLite.",
        "Implemented Azure Document Intelligence workflows with concurrent batches, retry handling, and visible per-file failure states.",
        "Extended guarded SAP and Excel consumption processing while keeping consequential actions under operator supervision.",
        "Built role-based operator views plus Databricks-backed assistance, deterministic local analytics, and native SVG charts.",
    ]
    for text in bullets:
        add_bullet(document, text, bullet_id)

    add_section(document, "Projects")
    add_role_line(document, "Portfolio & Interface Lab", "React | TypeScript | Tailwind | Vite")
    add_bullet(document, "Designed and built this responsive editorial portfolio with two themes, accessible interaction patterns, and 12 live component demos.", bullet_id)
    add_bullet(document, "Implemented a retrieval-grounded portfolio assistant with a Databricks-to-OpenRouter provider cascade and a scripted offline fallback.", bullet_id)

    add_section(document, "Education")
    education = profile["education"][0]
    add_role_line(document, education["institution"], education["end"])
    paragraph = document.add_paragraph()
    set_spacing(paragraph, after=1)
    set_font(paragraph.add_run(education["degree"]), 9.5)

    add_section(document, "Technical Skills")
    skills = [
        ("Languages", "Python, TypeScript, JavaScript, SQL, HTML, CSS"),
        ("Frontend", "React, Tailwind CSS, Framer Motion, responsive UI, accessibility"),
        ("Backend & AI", "FastAPI, REST APIs, SQLite, Azure Document Intelligence, Databricks, RAG, LLM APIs"),
        ("Tools", "Git, GitHub, Vite, Vercel, testing and CI workflows"),
    ]
    for label, values in skills:
        paragraph = document.add_paragraph()
        set_spacing(paragraph, after=1, line=1.0)
        set_font(paragraph.add_run(f"{label}: "), 9.25, bold=True)
        set_font(paragraph.add_run(values), 9.25)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build())
